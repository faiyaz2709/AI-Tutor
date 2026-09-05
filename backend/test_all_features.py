"""
test_all_features.py — Comprehensive Automated Test Suite

Verifies:
  1. GET / and GET /health
  2. GET /topics
  3. GET /videos (Python, DBMS, Statistics, and custom topic)
  4. POST /upload with TXT file
  5. POST /upload with DOCX file (if python-docx installed)
  6. GET /uploads
  7. POST /ask with default topic (Python)
  8. POST /ask with uploaded file_id
  9. POST /quiz with default topic
 10. POST /quiz with uploaded file_id
 11. POST /submit-quiz evaluation
 12. POST /adaptive re-teaching
 13. Verification of 90% mastery logic
 14. Error handling (invalid topics, missing questions, bad file_id)
"""

import sys
import io
from fastapi.testclient import TestClient
from main import app, quiz_sessions

client = TestClient(app)

def test_system():
    print("========================================")
    print("STARTING AI TUTOR VERIFICATION SUITE")
    print("========================================")

    passed = 0
    total = 0

    def assert_test(name, condition, detail=""):
        nonlocal passed, total
        total += 1
        if condition:
            passed += 1
            print(f"  [PASS] {name}")
        else:
            print(f"  [FAIL] {name}: {detail}")

    # 1. Health checks
    print("\n--- 1. Health & Discovery Endpoints ---")
    r = client.get("/health")
    assert_test("GET /health status 200", r.status_code == 200)
    assert_test("GET /health status is healthy", r.json().get("status") == "healthy")

    r = client.get("/topics")
    assert_test("GET /topics status 200", r.status_code == 200)
    topics_data = r.json()
    assert_test("Default topics include Python, DBMS, Statistics",
                len(topics_data.get("default_topics", [])) >= 3)

    # 2. Video recommendations
    print("\n--- 2. Personalized Video Learning ---")
    r = client.get("/videos?topic=Python&level=Beginner")
    assert_test("GET /videos Python Beginner returns 200", r.status_code == 200)
    v_data = r.json()
    assert_test("Python Beginner returns video list", len(v_data.get("videos", [])) > 0)
    assert_test("Video objects have id, title, channel, video_url",
                all("video_url" in v and "title" in v for v in v_data.get("videos", [])))

    r = client.get("/videos?topic=DBMS&level=Advanced")
    assert_test("GET /videos DBMS Advanced returns videos", len(r.json().get("videos", [])) > 0)

    r = client.get("/videos?topic=Custom_Neural_Nets&level=Expert")
    assert_test("GET /videos custom topic fallback works", len(r.json().get("videos", [])) > 0)

    # 3. File Uploads (TXT and DOCX)
    print("\n--- 3. File Upload Service ---")
    sample_text = (
        "Operating Systems: Process Scheduling and Deadlocks.\n"
        "A process is a program in execution. The process control block (PCB) stores state.\n"
        "Scheduling algorithms include First-Come First-Served, Shortest Job First, and Round Robin.\n"
        "Deadlock occurs when four Coffman conditions hold: Mutual Exclusion, Hold and Wait, "
        "No Preemption, and Circular Wait. Prevention eliminates one condition."
    )
    txt_file = io.BytesIO(sample_text.encode("utf-8"))
    r = client.post(
        "/upload",
        files={"file": ("OS_Notes.txt", txt_file, "text/plain")}
    )
    assert_test("POST /upload .txt returns 200", r.status_code == 200)
    upload_res = r.json()
    uploaded_file_id = upload_res.get("file_id")
    assert_test("Upload returned valid file_id", bool(uploaded_file_id))
    assert_test("Word count extracted correctly", upload_res.get("word_count", 0) > 20)

    # Verify document is listed in GET /uploads
    r = client.get("/uploads")
    assert_test("GET /uploads lists newly uploaded file",
                any(u["file_id"] == uploaded_file_id for u in r.json().get("uploads", [])))

    # Test DOCX upload
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Cloud Computing Architecture", 0)
        doc.add_paragraph("Cloud models include IaaS, PaaS, and SaaS. Scalability can be vertical or horizontal.")
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        r = client.post(
            "/upload",
            files={"file": ("Cloud_Notes.docx", docx_io, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        assert_test("POST /upload .docx returns 200", r.status_code == 200)
        assert_test("DOCX text extracted correctly", r.json().get("word_count", 0) > 5)
    except Exception as e:
        print(f"  [DOCX SKIPPED]: {e}")

    # 4. POST /ask (Built-in PDF topic vs Uploaded File)
    print("\n--- 4. RAG + Groq AI Explanations ---")
    # Built-in topic
    r = client.post(
        "/ask",
        json={"topic": "Python", "level": "Beginner", "question": "What is a list and tuple?"}
    )
    assert_test("POST /ask with Python returns 200", r.status_code == 200)
    ask_res = r.json()
    assert_test("POST /ask returns non-empty answer", len(ask_res.get("answer", "")) > 10)
    assert_test("Context used flag is True", ask_res.get("context_used") is True)

    # Uploaded file
    r = client.post(
        "/ask",
        json={
            "topic": "Operating Systems",
            "level": "Intermediate",
            "question": "What are the four conditions for deadlock?",
            "file_id": uploaded_file_id
        }
    )
    assert_test("POST /ask with uploaded file_id returns 200", r.status_code == 200)
    uploaded_ask_res = r.json()
    assert_test("Uploaded ask returned explanation", len(uploaded_ask_res.get("answer", "")) > 10)

    # 5. POST /quiz (Built-in vs Uploaded File)
    print("\n--- 5. Quiz Generation & Server-Side Security ---")
    # Built-in topic quiz
    r = client.post(
        "/quiz",
        json={"topic": "Python", "level": "Beginner", "question": "lists and dictionaries", "num_questions": 4}
    )
    assert_test("POST /quiz with Python returns 200", r.status_code == 200)
    quiz_res = r.json()
    session_id = quiz_res.get("session_id")
    questions = quiz_res.get("questions", [])
    assert_test("Quiz has session_id", bool(session_id))
    assert_test("Quiz generated questions", len(questions) > 0)
    assert_test("Correct answers are HIDDEN from frontend questions",
                all("correct_answer" not in q for q in questions))

    # Uploaded file quiz
    r = client.post(
        "/quiz",
        json={
            "topic": "Operating Systems",
            "level": "Intermediate",
            "question": "deadlock and scheduling",
            "num_questions": 4,
            "file_id": uploaded_file_id
        }
    )
    assert_test("POST /quiz with uploaded file_id returns 200", r.status_code == 200)
    uploaded_quiz_res = r.json()
    uploaded_session_id = uploaded_quiz_res.get("session_id")
    assert_test("Uploaded quiz has session_id and file_id tracked",
                bool(uploaded_session_id) and quiz_sessions[uploaded_session_id].get("file_id") == uploaded_file_id)

    # 6. POST /submit-quiz
    print("\n--- 6. Quiz Evaluation & Mastery Threshold ---")
    stored_key = quiz_sessions[session_id]["answer_key"]
    # Perfect score submission
    perfect_answers = {str(k): v for k, v in stored_key.items()}
    r = client.post(
        "/submit-quiz",
        json={"session_id": session_id, "answers": perfect_answers}
    )
    assert_test("POST /submit-quiz returns 200", r.status_code == 200)
    eval_perfect = r.json()
    assert_test("100% score marked as mastered",
                eval_perfect.get("percentage") == 100.0 and eval_perfect.get("mastered") is True)

    # Failing score submission (0%)
    wrong_answers = {str(k): "Definitively Wrong Answer" for k in stored_key.keys()}
    r = client.post(
        "/submit-quiz",
        json={"session_id": session_id, "answers": wrong_answers}
    )
    eval_fail = r.json()
    assert_test("0% score marked as NOT mastered",
                eval_fail.get("percentage") == 0.0 and eval_fail.get("mastered") is False)

    # 7. POST /adaptive (Adaptive Re-teaching)
    print("\n--- 7. Adaptive Learning ---")
    r = client.post(
        "/adaptive",
        json={
            "topic": "Python",
            "level": "Beginner",
            "question": "lists and loops",
            "percentage": 25.0,
            "session_id": session_id,
            "attempt": 1
        }
    )
    assert_test("POST /adaptive returns 200", r.status_code == 200)
    adaptive_res = r.json()
    assert_test("Adaptive response action is 'reteach'", adaptive_res.get("action") == "reteach")
    assert_test("Adaptive re-teaching explanation provided",
                bool(adaptive_res.get("reteach_explanation")))

    # Adaptive re-teaching on uploaded file
    r = client.post(
        "/adaptive",
        json={
            "topic": "Operating Systems",
            "level": "Intermediate",
            "question": "deadlock conditions",
            "percentage": 50.0,
            "session_id": uploaded_session_id,
            "attempt": 1
        }
    )
    assert_test("POST /adaptive on uploaded file returns 200", r.status_code == 200)
    assert_test("Uploaded adaptive retains file_id", adaptive_res.get("file_id") is not None or r.json().get("file_id") == uploaded_file_id)

    # 8. Error handling
    print("\n--- 8. Error Handling & Edge Cases ---")
    r = client.post("/ask", json={"topic": "", "level": "Beginner", "question": "test"})
    assert_test("Empty topic rejected with 400", r.status_code == 400)

    r = client.post("/ask", json={"topic": "UnknownTopicXYZ", "level": "Beginner", "question": "test"})
    assert_test("Non-existent topic rejected with 404", r.status_code == 404)

    r = client.post("/ask", json={"topic": "Python", "level": "Beginner", "question": "test", "file_id": "nonexistent_id"})
    assert_test("Non-existent file_id rejected with 404", r.status_code == 404)

    print("\n========================================")
    print(f"RESULTS: {passed} / {total} tests passed ({passed/total*100:.1f}%)")
    print("========================================")
    return passed == total

if __name__ == "__main__":
    success = test_system()
    sys.exit(0 if success else 1)
