"""
upload_service.py — Safe file upload, validation, and text extraction

Responsibilities:
  - Validate file type (PDF, TXT, DOCX) and size
  - Sanitize filenames to prevent path traversal or injection
  - Extract text safely from uploaded documents
  - Maintain an in-memory session registry of uploaded materials
  - Provide helper functions to retrieve context for RAG and Quiz flows
"""

import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional
from PyPDF2 import PdfReader

try:
    import docx
except ImportError:
    docx = None

# Safe upload storage directory
UPLOAD_DIR = Path(__file__).resolve().parent.parent / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Configuration
MAX_FILE_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}

# In-memory document registry: { file_id: { ...metadata, "text": "..." } }
_uploaded_documents: Dict[str, dict] = {}


def sanitize_filename(filename: str) -> str:
    """
    Sanitizes a filename to prevent path traversal and unsafe characters.
    """
    # Remove any directory path components
    clean_name = os.path.basename(filename)
    # Replace any character that is not alphanumeric, dot, underscore, or hyphen
    clean_name = re.sub(r'[^a-zA-Z0-9._-]', '_', clean_name)
    # Ensure reasonable length
    if len(clean_name) > 100:
        base, ext = os.path.splitext(clean_name)
        clean_name = base[:90] + ext
    return clean_name or "uploaded_document"


def extract_text_from_file(file_path: Path, file_ext: str) -> str:
    """
    Extracts text based on the file extension.
    """
    text = ""
    ext = file_ext.lower()

    if ext == ".pdf":
        try:
            reader = PdfReader(str(file_path))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to read PDF content: {str(e)}")

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")

    elif ext == ".docx":
        if docx is None:
            raise ValueError("DOCX extraction is not available. Please install python-docx.")
        try:
            doc = docx.Document(str(file_path))
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            text = "\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX document: {str(e)}")

    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Allowed types: PDF, TXT, DOCX.")

    if not text.strip():
        raise ValueError("The uploaded document appears to be empty or contains no extractable text.")

    return text.strip()


async def process_uploaded_file(file) -> dict:
    """
    Validates, saves, extracts text, and registers an uploaded file.
    Takes a FastAPI UploadFile object.
    """
    original_name = file.filename or "uploaded_document"
    ext = Path(original_name).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file extension '{ext}'. Only .pdf, .txt, and .docx files are accepted.")

    # Read content to check size
    content = await file.read()
    if not content:
        raise ValueError("Uploaded file is empty.")

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File size exceeds the 15 MB limit. (Size: {len(content) / (1024*1024):.2f} MB)")

    # Generate safe unique ID and path
    file_id = str(uuid.uuid4())
    safe_name = sanitize_filename(original_name)
    saved_filename = f"{file_id[:8]}_{safe_name}"
    save_path = UPLOAD_DIR / saved_filename

    # Save to disk safely
    try:
        with open(save_path, "wb") as f:
            f.write(content)
    except Exception as e:
        raise ValueError(f"Failed to save file on server: {str(e)}")

    # Extract text
    extracted_text = extract_text_from_file(save_path, ext)

    # Friendly display title
    display_title = Path(safe_name).stem.replace("_", " ").title()

    doc_meta = {
        "file_id": file_id,
        "filename": safe_name,
        "title": display_title,
        "file_type": ext.lstrip(".").upper(),
        "char_count": len(extracted_text),
        "word_count": len(extracted_text.split()),
        "saved_path": str(save_path),
        "uploaded_at": datetime.utcnow().isoformat(),
        "text": extracted_text
    }

    _uploaded_documents[file_id] = doc_meta

    return {
        "file_id": file_id,
        "filename": safe_name,
        "title": display_title,
        "file_type": ext.lstrip(".").upper(),
        "char_count": len(extracted_text),
        "word_count": len(extracted_text.split()),
        "message": f"Successfully processed '{safe_name}'! You can now ask questions or generate a quiz."
    }


def get_uploaded_document(file_id: str) -> Optional[dict]:
    """
    Returns document metadata and extracted text by file_id.
    """
    return _uploaded_documents.get(file_id)


def get_uploaded_file_text(file_id: str) -> Optional[str]:
    """
    Returns the extracted text for an uploaded document, or None if not found.
    """
    doc = _uploaded_documents.get(file_id)
    return doc["text"] if doc else None


def list_uploaded_documents() -> list:
    """
    Returns a list of all currently active uploaded documents (excluding raw text).
    """
    return [
        {
            "file_id": d["file_id"],
            "filename": d["filename"],
            "title": d["title"],
            "file_type": d["file_type"],
            "char_count": d["char_count"],
            "word_count": d["word_count"],
            "uploaded_at": d["uploaded_at"]
        }
        for d in _uploaded_documents.values()
    ]
